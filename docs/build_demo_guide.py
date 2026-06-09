from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "demo-guide-assets"
OUTPUT = ROOT / "docs" / "sanfu-style-lab-demo-guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "151515"
MUTED = "555555"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_east_asia(run, font_name="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    set_east_asia(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_east_asia(run)
    return p


def add_body(doc, text, bold=False, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)
    set_east_asia(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        set_east_asia(run)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia(run)


def add_screenshot(doc, filename, caption):
    path = ASSET_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(6.35))
    add_caption(doc, caption)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_east_asia(r)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(INK)
    set_east_asia(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1.45, 5.05])
    for idx, (label, detail) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), label, bold=True, color=DARK_BLUE)
        set_cell_text(table.cell(idx, 1), detail)
        set_cell_shading(table.cell(idx, 0), LIGHT_BLUE)
    return table


def add_module_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1.45, 2.35, 2.70])
    headers = ["模块", "用户怎么操作", "对营销的意义"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, color=DARK_BLUE)
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), text)
    return table


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("SANFU Style Lab Demo Guide")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia(run)


def build():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("三福潮搭实验室 Demo 产品营销说明文档")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(INK)
    set_east_asia(r)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    sr = subtitle.add_run("面向产品营销团队：用于理解 demo 操作路径、模块价值、现场演示话术和后续包装方向")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia(sr)

    add_callout(
        doc,
        "一句话定位",
        "这个 demo 不是普通商品列表，而是把“场景、风格、性别、单品、AI 建议、封面生成”串成一次可演示的年轻化潮搭体验。",
    )

    add_key_value_table(
        doc,
        [
            ("目标用户", "学生、年轻白领、逛三福时想快速找到搭配灵感的人群。"),
            ("营销目标", "把单品陈列升级为可讲故事、可试搭、可分享的生活状态入口。"),
            ("核心卖点", "先选场景和风格，再看对应男女商品素材，最终生成一张可传播的三福穿搭封面。"),
            ("演示重点", "评委/营销同事优先看“套装”完成效果；其他品类用于展示商品素材库与后续扩展能力。"),
        ],
    )

    add_screenshot(doc, "01-home-entry.png", "图 1：首页以品牌和“今日搭配”入口建立第一屏心智。")

    add_heading(doc, "1. Demo 体验路径", 1)
    add_body(doc, "建议演示时按照下面 5 步走，营销同事能快速理解这是一个完整闭环，而不是几个独立页面。")
    add_bullets(
        doc,
        [
            "进入首页：强调三福 Style Lab 是“生活状态型”的穿搭实验入口。",
            "选择角色：先点女生或男生，系统立刻切换对应风格和商品素材。",
            "打开灵感筛选：按生活场景、风格表达、潮搭频道逐层筛选。",
            "点击商品/清单：查看单品素材和详情弹窗，说明未来可接入价格、库存、购买入口。",
            "生成封面：把当前搭配转成一张可分享、可种草的三福潮搭封面。",
        ],
    )

    add_screenshot(doc, "02-lab-before-gender.png", "图 2：实验室主界面，选择性别前先让用户看到角色入口和完整操作框架。")

    add_heading(doc, "2. 模块拆解：每一块怎么操作、有什么意义", 1)
    add_module_table(
        doc,
        [
            ("生活场景", "点击 Step 01，选择通勤、校园、Livehouse、宿舍、约会、Citywalk 等场景。", "把商品推荐从“货架逻辑”变成“生活情境逻辑”，适合做内容营销和门店主题陈列。"),
            ("风格表达", "选择性别后，Step 02 自动出现对应风格，例如女生甜酷/韩系，男生机能/Clean Fit。", "同一批商品可被不同风格语言重新包装，方便输出小红书/短视频文案。"),
            ("潮搭频道", "Step 03 中“套装”独占一行；上装、下装、鞋子、包包、服饰配件、饰品两列排布。", "“套装”承担完整视觉样例，单品频道承担素材库展示和扩展能力。"),
            ("试衣展示区", "中间区域展示纸娃娃角色、场景背景和当前搭配层数。", "让用户形成“我在试搭”的参与感，比静态商品列表更容易停留。"),
            ("AI 好搭子", "右侧提供快速灵感按钮，例如变甜、变酷、约会适配、一键搭完整套。", "将导购语言产品化，适合包装成“陪逛/陪搭”卖点。"),
            ("搭配清单", "点击清单里的部位，弹出对应商品详情；没有选择时会用同品类素材展示。", "为后续接商品详情、价格、尺码、跳转购买留出自然入口。"),
            ("封面生成", "点击底部“生成三福潮搭封面”，进入上传/预览/生成流程。", "把试搭结果变成可分享物料，利于社交传播和活动裂变。"),
        ],
    )

    add_heading(doc, "3. 灵感筛选与商品素材展示", 1)
    add_body(doc, "左侧筛选仓是 demo 的“决策引导区”。营销同事讲解时可以强调：用户不是从海量 SKU 里硬找，而是先告诉系统“今天是什么场景、想要什么感觉”。")
    add_bullets(
        doc,
        [
            "套装：优先展示真实完成效果，适合评委和业务方快速判断视觉方向。",
            "单品：目前以真实商品素材图展示，图片区域放大，标签弱化，让用户先看清商品本身。",
            "男女分类：分类结构保持一致，但分类内商品不同，方便后续统一运营后台和内容模板。",
        ],
    )
    add_screenshot(doc, "03-female-products-filter.png", "图 3：选择女生和上装后，商品卡直接展示真实素材图，标签变为低干扰辅助信息。")

    add_heading(doc, "4. 当前搭配清单与详情弹窗", 1)
    add_body(doc, "当前搭配清单把“试搭过程”变成可查看、可解释、可继续扩展的商品结构。")
    add_bullets(
        doc,
        [
            "用户可点击任意部位，如上装、下装、鞋子，查看当前选择或同类素材。",
            "弹窗先展示大图和商品名称，后续可以接入价格、材质、库存、尺码、购买按钮。",
            "对营销团队来说，这里是从“灵感体验”转向“商品转化”的关键桥。",
        ],
    )
    add_screenshot(doc, "04-outfit-detail-modal.png", "图 4：点击搭配清单部位后出现商品详情弹窗，适合作为商品转化入口。")

    add_heading(doc, "5. 男生饰品与服饰配件：体现素材库扩展能力", 1)
    add_body(doc, "男生侧不仅有上装、下装、鞋子，也已经接入饰品和服饰配件素材，证明 demo 可以承载不同品类的真实图片。")
    add_bullets(
        doc,
        [
            "服饰配件：墨镜口罩、透明框眼镜等归入“服饰配件”，强调造型完整度。",
            "饰品：银色骷髅项链、民族风彩绳手链归入“饰品”，适合高街、复古、Citywalk 内容表达。",
            "讲解时可以说：分类统一，内容因性别和风格变化，后续加素材即可扩充。",
        ],
    )
    add_screenshot(doc, "05-male-jewelry-products.png", "图 5：男生饰品频道展示新接入的项链和手链素材。")

    add_heading(doc, "6. 封面生成：把试搭结果变成传播素材", 1)
    add_body(doc, "封面生成页是 demo 的营销闭环。它把用户的搭配选择转成一张可保存、可分享、可继续加工为活动物料的视觉结果。")
    add_bullets(
        doc,
        [
            "对用户：获得一张“今天的三福穿搭状态”封面，有参与感和分享动机。",
            "对门店：可作为活动打卡、社群晒单、会员任务的视觉入口。",
            "对品牌：把三福商品从单品售卖转成“生活方式内容”。",
        ],
    )
    add_screenshot(doc, "06-cover-generator.png", "图 6：封面生成入口承接试搭体验，适合做社交分享和活动玩法。")

    add_heading(doc, "7. 产品营销可以怎么讲", 1)
    add_callout(
        doc,
        "30 秒演示话术",
        "“这个 demo 把三福商品变成一次潮搭实验。用户先选今天的生活场景，再选性别和风格，系统展示对应商品素材；如果想快速出效果，可以看套装或让 AI 一键搭完整套。最后生成一张三福潮搭封面，既能作为个人分享，也能变成门店活动和内容种草素材。”",
    )
    add_body(doc, "对外沟通时建议突出三个关键词：")
    add_bullets(
        doc,
        [
            "场景化：不是卖一件衣服，而是卖“今天怎么穿”。",
            "年轻化：风格词、AI 好搭子、封面生成都更贴近年轻用户表达。",
            "可扩展：分类、素材、AI 文案、商品详情、购买入口都可以继续接入真实业务数据。",
        ],
    )

    add_heading(doc, "8. 后续可包装方向", 1)
    add_module_table(
        doc,
        [
            ("门店活动", "到店扫码选择场景，生成三福潮搭封面。", "提升互动感，形成可传播的门店打卡内容。"),
            ("会员运营", "会员上传照片或选择角色，领取专属搭配建议。", "增强会员粘性，沉淀偏好数据。"),
            ("内容种草", "按校园、通勤、Livehouse 等场景导出搭配主题。", "方便短视频、小红书、公众号做系列化内容。"),
            ("商品转化", "详情弹窗接入价格、库存、尺码、购买链接。", "把灵感浏览自然导向交易闭环。"),
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
